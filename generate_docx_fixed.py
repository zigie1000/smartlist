import sys
from docx import Document
from docx.shared import Inches
import os

# Usage: python3 generate_docx.py input.txt logo.png [image1.png image2.png ...]

input_txt = sys.argv[1]
logo_path = sys.argv[2]
image_paths = sys.argv[3:]
output_docx = "/tmp/PromptAgentHQ_Listing.docx"

doc = Document()

# Insert logo if available
if os.path.exists(logo_path):
    doc.add_picture(logo_path, width=Inches(2))
    doc.add_paragraph("")  # Spacer

# Insert text content
with open(input_txt, "r", encoding="utf-8") as f:
    text = f.read().strip()
    doc.add_paragraph(text)
    doc.add_paragraph("")  # Spacer

# Insert property images if any
for image_path in image_paths:
    if os.path.exists(image_path):
        doc.add_picture(image_path, width=Inches(5))
        doc.add_paragraph("")  # Spacer

# Save the document
doc.save(output_docx)
