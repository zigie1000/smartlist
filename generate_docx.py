
import sys
import os
import io
import base64
import tempfile
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def generate_docx(input_file, logo_b64=None, image_b64_list=[]):
    doc = Document()

    # Insert logo if provided
    if logo_b64:
        try:
            logo_stream = io.BytesIO(base64.b64decode(logo_b64.split(",")[1]))
            doc.add_picture(logo_stream, width=Inches(2.5))
            doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            doc.add_paragraph()
        except Exception as e:
            print("Logo error:", e)

    # Add the listing text
    with open(input_file, 'r', encoding='utf-8') as f:
        text = f.read()
    doc.add_paragraph(text)
    doc.add_paragraph()

    # Insert up to 3 images
    for img_b64 in image_b64_list[:3]:
        try:
            img_stream = io.BytesIO(base64.b64decode(img_b64.split(",")[1]))
            doc.add_picture(img_stream, width=Inches(5))
            doc.add_paragraph()
        except Exception as e:
            print("Image error:", e)

    return doc

if __name__ == "__main__":
    args = sys.argv[1:]
    input_file = args[0]
    logo_b64 = args[1] if len(args) > 1 else None
    images_b64 = args[2:] if len(args) > 2 else []
    doc = generate_docx(input_file, logo_b64, images_b64)
    output_docx = input_file.replace(".txt", ".docx")
    doc.save(output_docx)
